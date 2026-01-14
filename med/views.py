from urllib import request
from django.shortcuts import render
from django.shortcuts import render
from django.http import JsonResponse,HttpResponse
from .models import Patient, Report
from django.views.decorators.csrf import csrf_exempt
import json
from datetime import datetime
import base64
from docx import Document
from docx.shared import Inches
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# Create your views here.


def login(request):
    return render(request, 'login.html')

def profile(request):
    return render(request, 'profile.html')

def signup(request):
    return render(request,'signup.html')

def popupform(request):
    return render(request,'popupform.html')
# def report(request, id):
#     return render(request, 'report.html', {'id': id})

from django.shortcuts import render, get_object_or_404
from .models import UserAccount

from django.shortcuts import render, get_object_or_404, redirect
from .models import UserAccount

def user_detail(request, id):
    user = get_object_or_404(UserAccount, id=id)

    if request.method == "POST":
        # Login Details
        user.is_active = request.POST.get('status') == 'ACTIVE'
        # user.modality = request.POST.get('modality', user.modality)

        # Personal details
        user.name = request.POST.get('personName', user.name)
        # user.contact = request.POST.get('mobileNo', user.contact)
        # user.email = request.POST.get('email', user.email)

        # Password update
        new_password = request.POST.get('password')
        if new_password:
            user.password = new_password  # Optional: hash it in production

        user.save()
        return redirect('signup')  # Refresh page after saving

    return render(request, 'user_detail.html', {'user': user})


# views.py
from django.shortcuts import redirect

def update_user(request, id):
    user = get_object_or_404(UserAccount, id=id)
    if request.method == 'POST':
        user.name = request.POST.get('accountName')
        # user.contact = request.POST.get('mobileNo')
        user.usertype = request.POST.get('userType')
        user.is_active = request.POST.get('status') == 'ACTIVE'
        # add other fields
        user.save()
        return redirect('user_detail', id=user.id)
    return render(request, 'user_detail.html', {'user': user})


# def imagingA(request):
#     patients = Patient.objects.all().order_by('-entry_time')
#     return render(request, 'imagingA.html', {'patients': patients})

# def RADS(request):
#     patients = Patient.objects.all().order_by('-entry_time')
#     return render(request, 'RADS.html', {'patients': patients})

def invoice(request):
    return render(request, 'invoice.html')
def payment(request):
    return render(request, 'payment.html')


# @csrf_exempt
# def add_patient(request):
#     if request.method == 'POST':
#         try:
#             data = json.loads(request.body)
#             print("Received data:", data)  # For debugging
            
#             # Clean up base64 image data if needed
#             scan_image = data.get('scan', '')  # Use get() with default value
#             if ',' in scan_image:
#                 scan_image = scan_image.split(',')[1]
#             center_name = data.get('center') or request.session.get('user_name')  # ✅ from frontend or session

            
#             patient = Patient.objects.create(
#                 name=data.get('name'),
#                 age=data.get('age'),
#                 gender=data.get('gender'),
#                 history=data.get('history'),
#                 scan_type=data.get('scanType'),  # Changed from scan_type
#                 body_part=data.get('bodyPart'),  # Changed from body_part
#                 ref_by=data.get('refBy'),        # Changed from ref_by
#                 scan_image=scan_image,
#                 center=center_name  # ✅ store the center name
#             )
#             return JsonResponse({
#                 'status': 'success',
#                 'patient_id': patient.patient_id
#             })
#         except KeyError as e:
#             print("KeyError:", e)  # Add debug print
#             return JsonResponse({
#                 'status': 'error',
#                 'message': f'Missing field: {str(e)}'
#             })
#         except Exception as e:
#             print("Exception:", e)  # Add debug print
#             return JsonResponse({
#                 'status': 'error',
#                 'message': str(e)
#             })
@csrf_exempt
def add_patient(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            import json as json_module

            manual_id = data.get('patient_id')
            scan_data = data.get('scan', '')

            # Handle multiple images (array) or single image (backward compatibility)
            if isinstance(scan_data, list):
                # Multiple images - store as JSON array
                images_list = []
                for img in scan_data:
                    if isinstance(img, str):
                        # Remove data:image prefix if present
                        if ',' in img:
                            images_list.append(img.split(',')[1])
                        else:
                            images_list.append(img)
                    else:
                        images_list.append(str(img))
                scan_image = json_module.dumps(images_list)
            else:
                # Single image (old format - backward compatibility)
                scan_image = scan_data
                if isinstance(scan_image, str) and ',' in scan_image:
                    scan_image = scan_image.split(',')[1]
                # Convert to array format for consistency
                scan_image = json_module.dumps([scan_image]) if scan_image else json_module.dumps([])

            hospital_id = data.get('hospital_id')

            # 🏥 Hospital selection logic
            if hospital_id:
                hospital = UserAccount.objects.get(id=hospital_id, usertype='IMAGING')
                center_name = hospital.name
                created_by = hospital
            else:
                center_name = request.session.get('user_name')
                user_id = request.session.get('user_id')
                if user_id:
                    created_by = UserAccount.objects.get(id=user_id)
                else:
                    created_by = None

            # Patient ID
            if manual_id:
                patient_id = manual_id
            else:
                now = datetime.now()
                import random
                patient_id = f"P{now.strftime('%Y%m%d')}-{random.randint(1000,9999)}"

            patient = Patient.objects.create(
                patient_id=patient_id,
                name=data.get('name'),
                age=data.get('age'),
                gender=data.get('gender'),
                history=data.get('history'),
                scan_type=data.get('scanType'),
                body_part=data.get('bodyPart'),
                ref_by=data.get('refBy'),
                scan_image=scan_image,  # JSON array of images
                center=center_name,
                created_by=created_by   # ⭐ hospital user linked
            )

            return JsonResponse({'status': 'success', 'patient_id': patient.patient_id})

        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})



from django.shortcuts import render
from .models import Patient,UserAccount
def index(request):
    user_id = request.session.get('user_id')
    user_name = request.session.get('user_name')
    user_type = request.session.get('user_type')

    if not user_id:
        return redirect('login')

    # SUPERADMIN — All data
    if user_type == "SUPERADMIN":
        centers = UserAccount.objects.filter(usertype='IMAGING', is_active=True)
        rads = UserAccount.objects.filter(usertype='RADS', is_active=True)
        patients = Patient.objects.all().order_by('-entry_time')

    # ADMIN — Only his centers, rads and patients
    elif user_type == "ADMIN":
        centers = UserAccount.objects.filter(
            parent_admin_id=user_id,
            usertype='IMAGING',
            is_active=True
        )

        patients = Patient.objects.filter(
            created_by__in=centers
        ).order_by('-entry_time')

        rads = UserAccount.objects.filter(
            parent_admin_id=user_id,
            usertype='RADS',
            is_active=True
        )

    # IMAGING — Only own created patients
    elif user_type == "IMAGING":
        centers = None
        rads = None
        patients = Patient.objects.filter(
            created_by_id=user_id
        ).order_by('-entry_time')

    # RADS — Only assigned patients
    elif user_type == "RADS":
        centers = None
        rads = None
        patients = Patient.objects.filter(
            assigned_to_id=user_id
        ).order_by('-assigned_time','-entry_time')

    return render(request, 'index.html', {
        "user_name": user_name,
        "user_type": user_type,
        "centers": centers,
        "rads": rads,
        "patients": patients
    })


    # ✅ sab unique center names nikal lo (jo blank na ho)
    # centers_qs = Patient.objects.values_list('center', flat=True).distinct().order_by('center')
    # centers = [c for c in centers_qs if c and c.strip()]

    # ✅ dono bhej do template ko
   

@csrf_exempt
# from django.utils import timezone

def update_report(request, id):
    if request.method == 'POST':
        data = json.loads(request.body)
        try:
            patient = Patient.objects.get(id=id)

            patient.report = data.get('report', patient.report)
            patient.status = data.get('status', patient.status)

            tat_text = None

            # 🔥 TAT calculate only on FINAL
            if patient.status == 'FINAL':

                # Final time set only once
                if not patient.final_time:
                    patient.final_time = timezone.now()

                if patient.assigned_time:
                    delta = patient.final_time - patient.assigned_time
                    total_minutes = max(1, int(delta.total_seconds() // 60))

                    hours = total_minutes // 60
                    minutes = total_minutes % 60

                    tat_text = f"{hours}h {minutes}m" if hours else f"{minutes} min"

                    patient.tat = tat_text   # 🔥 SAVE TAT
            else:
                # 🔥 agar FINAL nahi hai to TAT clear
                patient.tat = None
                patient.final_time = None

            patient.save()

            return JsonResponse({
                'status': 'success',
                'tat': patient.tat,   # 🔥 always from DB
                'redirect': '/RADS/'
            })

        except Patient.DoesNotExist:
            return JsonResponse({
                'status': 'error',
                'message': 'Patient not found'
            })



def report(request, id):
    try:
        patient = Patient.objects.get(id=id)
        return render(request, 'report.html', {'patient': patient})
    except Patient.DoesNotExist:
        return render(request, 'report.html', {'error': 'Patient not found'})
    

@csrf_exempt
def get_reports(request, patient_id):
    try:
        patient = Patient.objects.get(id=patient_id)
        reports = patient.reports.all().values('id', 'report_text', 'status', 'created_at', 'updated_at')
        return JsonResponse({'reports': list(reports)})
    except Patient.DoesNotExist:
        return JsonResponse({'error': 'Patient not found'}, status=404)

@csrf_exempt
def create_report(request, patient_id):
    if request.method == 'POST':
        data = json.loads(request.body)
        try:
            patient = Patient.objects.get(id=patient_id)
            report_status = data.get('status', 'DRAFT')
            report = Report.objects.create(
                patient=patient,
                report_text=data.get('report_text', ''),
                status=report_status,
                created_by=UserAccount.objects.get(id=request.session.get('user_id')) if request.session.get('user_id') else None
            )
            
            # Update Patient status when report is FINAL
            if report_status == 'FINAL':
                patient.status = 'FINAL'
                
                # Set final_time only once
                if not patient.final_time:
                    patient.final_time = timezone.now()
                
                # Calculate TAT if assigned_time exists
                if patient.assigned_time:
                    delta = patient.final_time - patient.assigned_time
                    total_minutes = max(1, int(delta.total_seconds() // 60))
                    hours = total_minutes // 60
                    minutes = total_minutes % 60
                    tat_text = f"{hours}h {minutes}m" if hours else f"{minutes} min"
                    patient.tat = tat_text
                
                # Update patient report text from the report
                patient.report = report.report_text
                patient.save()
            
            return JsonResponse({'id': report.id, 'message': 'Report created'})
        except Patient.DoesNotExist:
            return JsonResponse({'error': 'Patient not found'}, status=404)

@csrf_exempt
def get_report(request, report_id):
    try:
        report = Report.objects.get(id=report_id)
        data = {
            'id': report.id,
            'report_text': report.report_text,
            'status': report.status,
            'created_at': report.created_at,
            'updated_at': report.updated_at
        }
        return JsonResponse(data)
    except Report.DoesNotExist:
        return JsonResponse({'error': 'Report not found'}, status=404)

@csrf_exempt
def update_report(request, report_id):
    if request.method == 'POST':
        data = json.loads(request.body)
        try:
            report = Report.objects.get(id=report_id)
            report.report_text = data.get('report_text', report.report_text)
            new_status = data.get('status', report.status)
            report.status = new_status
            report.save()
            
            # Update Patient status when report is FINAL
            patient = report.patient
            if new_status == 'FINAL':
                patient.status = 'FINAL'
                
                # Set final_time only once
                if not patient.final_time:
                    patient.final_time = timezone.now()
                
                # Calculate TAT if assigned_time exists
                if patient.assigned_time:
                    delta = patient.final_time - patient.assigned_time
                    total_minutes = max(1, int(delta.total_seconds() // 60))
                    hours = total_minutes // 60
                    minutes = total_minutes % 60
                    tat_text = f"{hours}h {minutes}m" if hours else f"{minutes} min"
                    patient.tat = tat_text
                
                # Update patient report text from the report
                patient.report = report.report_text
            else:
                # If status is DRAFT, keep patient status as is (don't change to FINAL)
                # Only clear TAT and final_time if needed
                if patient.status == 'FINAL':
                    # Don't change patient status if it's already FINAL
                    pass
                else:
                    patient.tat = None
                    patient.final_time = None
            
            patient.save()
            
            return JsonResponse({'message': 'Report updated', 'status': 'success'})
        except Report.DoesNotExist:
            return JsonResponse({'error': 'Report not found'}, status=404)

@csrf_exempt
def delete_report(request, report_id):
    if request.method == 'DELETE':
        try:
            report = Report.objects.get(id=report_id)
            report.delete()
            return JsonResponse({'message': 'Report deleted'})
        except Report.DoesNotExist:
            return JsonResponse({'error': 'Report not found'}, status=404)

    
@csrf_exempt
def get_patient(request, id):
    try:
        patient = Patient.objects.get(id=id)
        data = {
            'id': patient.id,
            'patient_id': patient.patient_id,
            'name': patient.name,
            'age': patient.age,
            'gender': patient.gender,
            'history': patient.history,
            'scan_type': patient.scan_type,
            'body_part': patient.body_part,
            'ref_by': patient.ref_by,
            'scan_image': patient.scan_image,
            'status': patient.status,
            'report': patient.report
        }
        return JsonResponse(data)
    except Patient.DoesNotExist:
        return JsonResponse({'error': 'Patient not found'}, status=404)





from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from .models import Patient  # apne model ka naam use karo

# @csrf_exempt
# def delete_patient(request, id):
#     if request.method == 'DELETE':
#         try:
#             patient = Patient.objects.get(id=id)
#             patient.delete()
#             return JsonResponse({'success': True})
#         except Patient.DoesNotExist:
#             return JsonResponse({'error': 'Patient not found'}, status=404)
#     return JsonResponse({'error': 'Invalid request'}, status=400)


        
        
@csrf_exempt
def download_report(request, id, format):
    try:
        patient = Patient.objects.get(id=id)
        if patient.status != 'FINAL':
            return JsonResponse({'error': 'Report not finalized'}, status=400)

        if format == 'pdf':
            # Create PDF
            buffer = BytesIO()
            p = canvas.Canvas(buffer, pagesize=letter)
            
            # Add content
            p.setFont("Helvetica", 12)
            p.drawString(100, 750, f"Patient ID: {patient.patient_id}")
            p.drawString(100, 730, f"Name: {patient.name}")
            p.drawString(100, 710, f"Age/Gender: {patient.age}/{patient.gender}")
            p.drawString(100, 690, f"Scan Type: {patient.scan_type}")
            p.drawString(100, 670, f"Body Part: {patient.body_part}")
            p.drawString(100, 650, f"Referred By: {patient.ref_by}")
            p.drawString(100, 630, "Report:")
            p.drawString(120, 610, patient.report or "No report available")
            
            p.save()
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="patient_report_{patient.patient_id}.pdf"'
            return response

        elif format == 'docx':
            # Create Word document
            doc = Document()
            doc.add_heading('Patient Report', 0)
            
            # Add content
            doc.add_paragraph(f'Patient ID: {patient.patient_id}')
            doc.add_paragraph(f'Name: {patient.name}')
            doc.add_paragraph(f'Age/Gender: {patient.age}/{patient.gender}')
            doc.add_paragraph(f'Scan Type: {patient.scan_type}')
            doc.add_paragraph(f'Body Part: {patient.body_part}')
            doc.add_paragraph(f'Referred By: {patient.ref_by}')
            doc.add_heading('Report:', level=1)
            doc.add_paragraph(patient.report or "No report available")
            
            # Save document
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="patient_report_{patient.patient_id}.docx"'
            return response

        return JsonResponse({'error': 'Invalid format'}, status=400)
        
    except Patient.DoesNotExist:
        return JsonResponse({'error': 'Patient not found'}, status=404)
    
    
from django.shortcuts import render, redirect
from .models import UserAccount
from .forms import SignupForm, LoginForm
def signup_view(request):
    session_user_id = request.session.get('user_id')
    session_user_type = request.session.get('user_type')

    if request.method == 'POST':
        form = SignupForm(request.POST)
        if form.is_valid():
            new_user = form.save(commit=False)
            if session_user_id and session_user_type == "ADMIN":
                try:
                    parent_admin = UserAccount.objects.get(id=session_user_id, usertype='ADMIN')
                    if new_user.usertype in ['IMAGING', 'RADS']:
                        new_user.parent_admin = parent_admin
                except UserAccount.DoesNotExist:
                    pass
            new_user.save()
            messages.success(request, f"User '{new_user.name}' created successfully.")
            if session_user_type == 'ADMIN':
                return redirect('user_list')
            elif session_user_type == 'SUPERADMIN':
                return redirect('super_admin')
            else:
                return redirect('login')
        else:
            messages.error(request, "Please fix errors.")
    else:
        form = SignupForm()

    # Important: provide users to template so admin sees own child users on signup page
    if session_user_type == "SUPERADMIN":
        users = UserAccount.objects.filter(usertype__in=['ADMIN','IMAGING','RADS']).order_by('-id')
    elif session_user_type == "ADMIN":
        users = UserAccount.objects.filter(parent_admin_id=session_user_id).order_by('-id')
    else:
        users = UserAccount.objects.none()

    admins = UserAccount.objects.filter(usertype='ADMIN', is_active=True)
    return render(request, 'signup.html', {'form': form, 'admins': admins, 'users': users})



from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from .models import UserAccount
from .forms import SignupForm, LoginForm
from django.contrib.auth.hashers import check_password

def login_view(request):
    form = LoginForm(request.POST or None)
    error = None

    if form.is_valid():
        userid = form.cleaned_data['userid']
        password = form.cleaned_data['password']

        try:
            user = UserAccount.objects.get(userid=userid)
            db_pass = user.password

            # 🔥 CASE 1: password HASHED है (new users)
            if db_pass.startswith("pbkdf2_"):
                if not check_password(password, db_pass):
                    raise UserAccount.DoesNotExist

            # 🔥 CASE 2: password PLAIN TEXT है (old users)
            else:
                if password != db_pass:
                    raise UserAccount.DoesNotExist

            # 🟢 Login successful → Session set
            request.session['user_id'] = user.id
            request.session['user_name'] = user.name
            request.session['user_type'] = user.usertype  

            # 🔀 Redirect based on user type
            if user.usertype == 'SUPERADMIN':
                return redirect('super_admin')
            elif user.usertype == 'ADMIN':
                return redirect('index')
            elif user.usertype == 'IMAGING':
                return redirect('imagingA')
            elif user.usertype == 'RADS':
                return redirect('RADS')

            return redirect('login')

        except UserAccount.DoesNotExist:
            error = "Invalid userid or password"

    return render(request, 'login.html', {'form': form, 'error': error})




def imagingA(request):
    if 'user_id' not in request.session:
        return redirect('login')

    user_id = request.session['user_id']
    user_name = request.session['user_name']

    # ✅ Admin ko sab dikhayenge
    if user_name.lower() == "admin":
        patients = Patient.objects.all().order_by('-entry_time')
    else:
        patients = Patient.objects.filter(created_by_id=user_id).order_by('-entry_time')

    return render(request, 'imagingA.html', {'patients': patients})

def RADS(request):
    if 'user_id' not in request.session:
        return redirect('login')

    user_id = request.session['user_id']
    user_name = request.session['user_name']

    # Agar admin ho toh sab dikhana
    if user_name.lower() == "admin":
        patients = Patient.objects.all().order_by('-entry_time')
    else:
        # Agar current logged-in user RAD type hai to assigned patients dikhao
        user = UserAccount.objects.get(id=user_id)
        if user.usertype == 'RADS':
            patients = Patient.objects.filter(assigned_to_id=user_id).order_by('-entry_time')
        else:
            # agar imaging center user hai to apne created patients dikhao (jaise pehle)
            patients = Patient.objects.filter(created_by_id=user_id).order_by('-entry_time')

    return render(request, 'RADS.html', {'patients': patients})


def logout_view(request):
    request.session.flush()  # session clear
    return redirect('login')

# def rads_page(request):
#     return render(request, 'RADS.html')

# def imaging_page(request):
#     return render(request, 'imagingA.html')




    # views.py (kisi jagah imports ke upar)
from django.http import JsonResponse, HttpResponseForbidden, HttpResponseBadRequest
from django.views.decorators.http import require_POST
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt  # agar aap csrf_token use nahi kar rahe toh rakh sakte ho

# Add this function somewhere in views.py
@csrf_exempt   # agar aap template me CSRF token bhejna nahi chahte toh isko rakh lo; production me hata ke proper CSRF use karo
@require_POST
def assign_patient(request, patient_id):
    print("Session check ->", request.session.get('user_name'), request.session.get('user_type'))

    # Session check
    if 'user_id' not in request.session:
        return HttpResponseForbidden("Login required")

    # ✔ Correct permission check
    user_type = request.session.get('user_type', '')

    # SUPERADMIN + ADMIN + IMAGING allowed
    if user_type not in ['SUPERADMIN', 'ADMIN', 'IMAGING']:
        return HttpResponseForbidden("Not allowed to assign")

    # get data
    try:
        data = json.loads(request.body.decode('utf-8'))
    except Exception:
        return HttpResponseBadRequest("Invalid JSON")

    rad_id = data.get('rad_id')

    # get patient
    try:
        patient = Patient.objects.get(id=patient_id)
    except Patient.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Patient not found'}, status=404)

    # Unassign
    # Unassign
    if not rad_id:
        patient.assigned_to = None
        patient.assigned_time = None

        # 🔥 ADD THESE LINES
        patient.status = 'UNREAD'
        patient.tat = None      # report unread hona chahiye
        patient.report = ''            # optional (agar report clear chahiye)
        patient.final_time = None      # optional (agar final reset chahiye)

        patient.save()

        return JsonResponse({
            'success': True,
            'assigned_to': None,
            'status': 'UNREAD'
        })


    # Assign to RADS
    try:
        rad_user = UserAccount.objects.get(id=rad_id, is_active=True, usertype='RADS')
    except UserAccount.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'RAD user not found'}, status=404)

    patient.assigned_to = rad_user
    patient.assigned_time = timezone.now()   # ⭐ yahi main fix
    patient.status = 'UNREAD'
    patient.save()

    return JsonResponse({
        'success': True,
        'assigned_to': {'id': rad_user.id, 'name': rad_user.name}
    })


from django.contrib import messages
from django.shortcuts import redirect
from django.contrib.auth.hashers import make_password  # optional, recommended
from .models import UserAccount
from datetime import date, datetime, timedelta
import json

def super_admin(request):

    # OLD DATA YOU ALREADY HAD
    users = UserAccount.objects.filter(usertype='ADMIN').order_by('-id')
    patients = Patient.objects.all().order_by('-entry_time')
    centers = UserAccount.objects.filter(usertype='IMAGING', is_active=True)
    admins = UserAccount.objects.filter(usertype='ADMIN', is_active=True)
    rads = UserAccount.objects.filter(usertype='RADS', is_active=True)

    # ⭐ NEW LIVE STATS (for KPI cards)
    total_patients = Patient.objects.count()
    total_centers = centers.count()
    total_rads = rads.count()
    total_admins = admins.count()

    # ⭐ NEW: Pending & Completed reports
    # ⭐ Pending = UNREAD + PENDING (if any)
    pending_count = Patient.objects.filter(status__iexact='UNREAD').count()
    pending_count += Patient.objects.filter(status__iexact='PENDING').count()

    # ⭐ Completed = FINAL
    completed_count = Patient.objects.filter(status__iexact='FINAL').count()



    # ⭐ NEW: Daily scans for last 7 days (line chart)
    daily_labels = []
    daily_scans = []

    for i in range(6, -1, -1):
        day = date.today() - timedelta(days=i)
        daily_labels.append(day.strftime("%b %d"))

        count = Patient.objects.filter(entry_time__date=day).count()
        daily_scans.append(count)

    return render(request, 'super_admin.html', {
        # original context
        'users': users,
        'patients': patients,
        'rads': rads,
        'centers': centers,
        'admins': admins,

        # ⭐ NEW values for Dashboard UI
        'total_patients': total_patients,
        'total_centers': total_centers,
        'total_rads': total_rads,
        'total_admins': total_admins,

        'pending_count': pending_count,
        'completed_count': completed_count,

        'daily_labels': json.dumps(daily_labels),
        'daily_scans': json.dumps(daily_scans),
    })


    

def toggle_user_status(request, id):
    user = UserAccount.objects.get(id=id)
    user.is_active = not user.is_active
    user.save()
    return redirect('super_admin')


def change_user_role(request, id, role):
    user = UserAccount.objects.get(id=id)
    if role in ['RADS', 'IMAGING']:
        user.usertype = role
        user.save()
    return redirect('super_admin')

from django.contrib import messages
from django.shortcuts import redirect
from .models import Patient

def assign_patient_superadmin(request):
    if request.method == "POST":
        patient_id = request.POST.get("patient_id")
        rads = request.POST.get("rads")

        try:
            patient = Patient.objects.get(id=patient_id)

            # ✅ Sirf RADS assign karo
            if rads:
                patient.assigned_to_id = rads
                # patient.status = "ASSIGNED"
                patient.save()

            messages.success(request, "Patient successfully assigned to RADS.")
        except Patient.DoesNotExist:
            messages.error(request, "Patient not found.")
        except Exception as e:
            messages.error(request, f"Assignment failed: {str(e)}")

    return redirect('super_admin')



from .models import Patient, UserAccount  # example model names

def patients(request):
    patients = Patient.objects.all().order_by('-id')
    rads = UserAccount.objects.filter(usertype='rads')  # or as per your model field
    return render(request, 'patients.html', {'patients': patients, 'rads': rads})

from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
from django.contrib import messages
from django.http import JsonResponse
from .models import UserAccount

@csrf_exempt
@require_POST
def add_user(request):
    name = request.POST.get('name')
    userid = request.POST.get('userid')
    password = request.POST.get('password')
    usertype = request.POST.get('usertype')
    parent_admin_id = request.POST.get('parent_admin_id')

    if not (name and userid and password and usertype):
        messages.error(request, "सभी फ़ील्ड भरें!")
        return redirect('super_admin')

    if UserAccount.objects.filter(userid=userid).exists():
        messages.error(request, "यह UserID पहले से मौजूद है!")
        return redirect('super_admin')

    user = UserAccount(name=name, userid=userid, password=password, usertype=usertype, is_active=True)

    if usertype in ['IMAGING', 'RADS'] and parent_admin_id:
        try:
            parent_admin = UserAccount.objects.get(id=parent_admin_id, usertype='ADMIN')
            user.parent_admin = parent_admin
        except UserAccount.DoesNotExist:
            messages.warning(request, "Admin नहीं मिला, user बिना parent admin के बना दिया गया।")

    user.save()
    messages.success(request, f"User '{name}' सफलतापूर्वक बना दिया गया ✅")
    return redirect('user_list')


from django.shortcuts import render, get_object_or_404
from .models import UserAccount

def admin_details(request, admin_id):
    admin_user = get_object_or_404(UserAccount, id=admin_id, usertype='ADMIN')

    imaging_users = UserAccount.objects.filter(usertype='IMAGING', parent_admin_id=admin_id)
    rads_users = UserAccount.objects.filter(usertype='RADS', parent_admin_id=admin_id)

    return render(request, 'admin_details.html', {
        'admin': admin_user,
        'imaging_users': imaging_users,
        'rads_users': rads_users,
    })

from django.http import JsonResponse

def admin_details_api(request, admin_id):
    admin = get_object_or_404(UserAccount, id=admin_id, usertype='ADMIN')
    imaging = UserAccount.objects.filter(usertype='IMAGING', parent_admin_id=admin_id)
    rads = UserAccount.objects.filter(usertype='RADS', parent_admin_id=admin_id)

    return JsonResponse({
        'name': admin.name,
        'userid': admin.userid,
        'imaging_count': imaging.count(),
        'rads_count': rads.count(),
        'imaging_list': list(imaging.values('name', 'userid', 'id')),
        'rads_list': list(rads.values('name', 'userid', 'id')),
    })



from django.shortcuts import render
from .models import UserAccount  # ya jo bhi model hai
from .forms import SignupForm

# def user_list(request):
#     user_id = request.session.get('user_id')
#     user_type = request.session.get('user_type')

#     if user_type == "SUPERADMIN":
#         users = UserAccount.objects.exclude(usertype='SUPERADMIN')

#     elif user_type == "ADMIN":
#         # सिर्फ उसी admin के imaging + rads users दिखाओ
#         users = UserAccount.objects.filter(parent_admin_id=user_id)

#     else:
#         messages.error(request, "ACCESS DENIED")
#         return redirect('index')

#     form = SignupForm()

#     return render(request, 'user_list.html', {
#         'form': form,
#         'users': users
#     })

def user_list(request):
    session_user_id = request.session.get('user_id')
    session_user_type = request.session.get('user_type')

    if not session_user_id:
        return redirect('login')

    if session_user_type == "SUPERADMIN":
        # show all ADMINs (or all users if you prefer)
        users = UserAccount.objects.filter(usertype='ADMIN').order_by('-id')
        admins = UserAccount.objects.filter(usertype='ADMIN', is_active=True)
    elif session_user_type == "ADMIN":
        # show IMAGING and RADS users whose parent_admin is this admin
        users = UserAccount.objects.filter(parent_admin_id=session_user_id).order_by('-id')
        admins = UserAccount.objects.filter(usertype='ADMIN')  # for select dropdown if needed
    else:
        # messages.error(request, "Access denied")
        return redirect('imagingA')

    form = SignupForm()
    return render(request, 'user_list.html', {'form': form, 'users': users, 'admins': admins})





from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from .models import UserAccount

def user_details_api(request, user_id):
    user = get_object_or_404(UserAccount, id=user_id)
    return JsonResponse({
        'name': user.name,
        'userid': user.userid,
        'usertype': user.usertype,
        'is_active': user.is_active,
        'status': 'Active' if user.is_active else 'Inactive',
        'password': user.password,  # ⚠️ सिर्फ development के लिए
        'created_at': user.created_at.strftime('%d %b %Y, %I:%M %p') if user.created_at else '-',
        'parent_admin': user.parent_admin.name if user.parent_admin else None,
        'parent_admin_id': user.parent_admin.id if user.parent_admin else None,
    })


@csrf_exempt
def update_user(request, user_id):
    """Update user details"""
    if request.method == 'POST':
        user = get_object_or_404(UserAccount, id=user_id)

        user.name = request.POST.get('name', user.name)
        user.userid = request.POST.get('userid', user.userid)
        user.usertype = request.POST.get('usertype', user.usertype)
        user.is_active = request.POST.get('is_active') == '1'

        # Update password only if provided
        new_password = request.POST.get('password')
        if new_password:
            user.password = new_password

        # Update parent admin
        parent_admin_id = request.POST.get('parent_admin_id')
        if parent_admin_id:
            try:
                parent_admin = UserAccount.objects.get(id=parent_admin_id, usertype='ADMIN')
                user.parent_admin = parent_admin
            except UserAccount.DoesNotExist:
                pass
        else:
            user.parent_admin = None

        user.save()
        messages.success(request, f"User '{user.name}' successfully updated! ✅")
        return redirect('super_admin')

    return redirect('user_list')


def delete_user(request, user_id):
    """Delete user"""
    user = get_object_or_404(UserAccount, id=user_id)
    user_name = user.name

    # Check if user has any associated patients
    patients_count = Patient.objects.filter(created_by=user).count()
    assigned_patients_count = Patient.objects.filter(assigned_to=user).count()

    if patients_count > 0 or assigned_patients_count > 0:
        messages.warning(request, f"Cannot delete user '{user_name}' - they have {patients_count + assigned_patients_count} associated patients. Please reassign or delete patients first.")
    else:
        user.delete()
        messages.success(request, f"User '{user_name}' successfully deleted! 🗑️")

    return redirect('user_list')

from .models import UserAccount, SuperAdminCreatedUsers
from django.contrib.auth.hashers import make_password
from django.views.decorators.http import require_POST

@require_POST
def create_user_page(request):
    session_user_id = request.session.get('user_id')
    session_user_type = request.session.get('user_type')

    if not session_user_id:
        return redirect('login')

    name = request.POST.get('name')
    userid = request.POST.get('userid')
    password = request.POST.get('password')
    usertype = request.POST.get('usertype')
    parent_admin = None

    # Duplicate check
    if UserAccount.objects.filter(userid=userid).exists():
        messages.error(request, "USER ID already exists.")
        return redirect('user_list')

    # If current logged-in user is ADMIN -> auto set parent_admin
    if session_user_type == "ADMIN":
        try:
            parent_admin = UserAccount.objects.get(id=session_user_id, usertype='ADMIN')
        except UserAccount.DoesNotExist:
            parent_admin = None

    # If SUPERADMIN -> optional parent_admin chosen from form
    elif session_user_type == "SUPERADMIN":
        parent_id = request.POST.get('parent_admin_id')
        if parent_id:
            try:
                parent_admin = UserAccount.objects.get(id=parent_id, usertype='ADMIN')
            except UserAccount.DoesNotExist:
                parent_admin = None

    # Create new user (hash password)
    new_user = UserAccount.objects.create(
        name=name,
        userid=userid,
        password=make_password(password),
        usertype=usertype.upper(),   # ensure uppercase consistency
        is_active=True,
        parent_admin=parent_admin
    )

    # optional logging table
    if session_user_type == "SUPERADMIN" and new_user.usertype in ['IMAGING', 'RADS', 'ADMIN']:
        SuperAdminCreatedUsers.objects.create(
            name=new_user.name,
            userid=new_user.userid,
            usertype=new_user.usertype
        )

    messages.success(request, f"{usertype} user '{name}' created.")

    # Redirect back based on who created
    if session_user_type == "ADMIN":
        return redirect('signup')        # admin should see its own user_list
    elif session_user_type == "SUPERADMIN":
        return redirect('super_admin')
    else:
        return redirect('login')




def admin_details_page(request, admin_id):
    admin_user = get_object_or_404(UserAccount, id=admin_id, usertype='ADMIN')

    imaging_users = UserAccount.objects.filter(
        usertype='IMAGING',
        parent_admin_id=admin_id
    )

    rads_users = UserAccount.objects.filter(
        usertype='RADS',
        parent_admin_id=admin_id
    )

    return render(request, 'admin_details_page.html', {
        'admin': admin_user,
        'imaging_users': imaging_users,
        'rads_users': rads_users
    })

from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
import json
from .models import Patient

@csrf_exempt
def delete_multiple(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            ids = data.get('ids', [])
            Patient.objects.filter(id__in=ids).delete()
            return JsonResponse({'success': True})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid request'}, status=400)


@csrf_exempt
def edit_patient(request, id):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            p = Patient.objects.get(id=id)
            p.name = data.get('name', p.name)
            p.age = data.get('age', p.age)
            p.body_part = data.get('body_part', p.body_part)
            p.history = data.get('history', p.history)
            p.ref_by = data.get('ref_by', p.ref_by)
            p.save()
            return JsonResponse({'success': True})
        except Patient.DoesNotExist:
            return JsonResponse({'success': False, 'error': 'Patient not found'}, status=404)
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid request'}, status=400)

from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
from django.contrib import messages
from django.shortcuts import redirect, get_object_or_404
from .models import UserAccount

MASTER_PASSWORD = "AXRIX@2024"

@csrf_exempt
@require_POST
def impersonate_with_password(request, user_id):

    # SUPERADMIN check (session based)
    if request.session.get('user_type') != "SUPERADMIN":
        messages.error(request, "Permission denied.")
        return redirect('user_list')

    # Master password verify
    password = request.POST.get("master_password")
    if password != MASTER_PASSWORD:
        messages.error(request, "Invalid master password.")
        return redirect('user_list')

    # Target user fetch
    target_user = get_object_or_404(UserAccount, id=user_id)

    # Clear existing login
    request.session.flush()

    # Set new impersonation session
    request.session['user_id'] = target_user.id
    request.session['user_name'] = target_user.name
    request.session['user_type'] = target_user.usertype

    messages.success(request, f"Logged in as {target_user.name}")

    # Redirect by usertype
    if target_user.usertype == "ADMIN":
        return redirect('index')

    if target_user.usertype == "IMAGING":
        return redirect('imagingA')

    if target_user.usertype == "RADS":
        return redirect('RADS')

    # Default
    return redirect('index')


import base64
from django.http import JsonResponse
from .models import Patient
import json

@csrf_exempt
def save_cropped_image(request, pk):
    if request.method == "POST":
        try:
            import json as json_module
            patient = Patient.objects.get(id=pk)
            data = json.loads(request.body)
            cropped_img = data.get("scan_image")
            image_index = data.get("image_index", 0)  # Which image to update

            # Remove base64 header
            if ',' in cropped_img:
                cropped_img = cropped_img.split(",")[1]

            # Get existing images
            images_list = patient.get_images()
            
            # Update specific image or add new
            if image_index < len(images_list):
                images_list[image_index] = cropped_img
            else:
                images_list.append(cropped_img)
            
            # Save back as JSON
            patient.set_images(images_list)
            patient.save()

            return JsonResponse({"status": "success"})

        except Patient.DoesNotExist:
            return JsonResponse({"status": "error", "message": "Patient not found"})


from django.http import JsonResponse
from .models import UserAccount

def admin_hospitals_api(request):
    print("SESSION =>", dict(request.session))

    user_id = request.session.get('user_id')
    user_type = request.session.get('user_type')

    print("ADMIN ID:", user_id)
    print("USER TYPE:", user_type)

    hospitals = UserAccount.objects.filter(
        parent_admin_id=user_id,
        usertype='IMAGING',
        is_active=True
    )

    print("HOSPITAL COUNT:", hospitals.count())

    return JsonResponse({
        'status': 'success',
        'hospitals': list(hospitals.values('id','name'))
    })

